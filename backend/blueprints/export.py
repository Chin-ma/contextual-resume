# backend/blueprints/export.py

from flask import Blueprint, jsonify, send_file, request, current_app 
import io
from docx import Document
from docx.shared import Pt # Only Pt is still relevant for docx font sizes
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import ReportLab components for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import inch 

export_bp = Blueprint('export', __name__)

@export_bp.route('/', methods=['POST'])
@export_bp.route('', methods=['POST'])
def export_resume():
    data = request.get_json()
    export_format = data.get('format', 'docx').lower()

    original_extracted_data = data.get('original_extracted_data', {})
    improved_summary = data.get('improved_summary', '')
    improved_bullets = data.get('improved_bullets', [])
    suggested_skills = data.get('suggested_skills', [])
    match_analysis = data.get('match_analysis', '') # Not used for export to PDF but kept

    def get_section_text(section_key, default_text=""):
        return original_extracted_data.get(section_key, default_text)

    name = get_section_text('contact_info', {}).get('name', 'Applicant Name')
    email = get_section_text('contact_info', {}).get('email', '')
    phone = get_section_text('contact_info', {}).get('phone', '')
    linkedin = get_section_text('contact_info', {}).get('linkedin', '')
    
    education_entries = original_extracted_data.get('education', [])
    skills_list = original_extracted_data.get('skills', [])
    achievements_list = original_extracted_data.get('achievements', [])


    if export_format == 'docx':
        document = Document()
        
        # Set basic margins (optional) - for DOCX, using Inches from docx.shared is correct
        from docx.shared import Inches # Re-import Inches just for DOCX context if needed
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # 1. Contact Info (Header-like)
        header_para = document.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runner = header_para.add_run(name)
        runner.bold = True
        runner.font.size = Pt(20)
        
        contact_info_para = document.add_paragraph()
        contact_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = []
        if email: contact_parts.append(email)
        if phone: contact_parts.append(phone)
        if linkedin: contact_parts.append(linkedin)
        contact_info_para.add_run(" | ".join(contact_parts))
        
        document.add_paragraph().add_run().add_break()


        # 2. Summary
        if improved_summary:
            document.add_heading('Summary', level=2)
            document.add_paragraph(improved_summary)
            document.add_paragraph().add_run().add_break()

        # 3. Experience
        if improved_bullets:
            document.add_heading('Experience', level=2)
            for bullet in improved_bullets:
                p = document.add_paragraph(style='List Bullet')
                p.add_run(bullet.lstrip('* ').strip())
            document.add_paragraph().add_run().add_break()


        # 4. Education (from original extracted data)
        if education_entries:
            document.add_heading('Education', level=2)
            for edu in education_entries:
                edu_para = document.add_paragraph()
                edu_para.add_run(f"{edu.get('degree', '')}").bold = True
                edu_para.add_run(f", {edu.get('university', '')}")
                if edu.get('year'):
                    edu_para.add_run(f", {edu.get('year', '')}")
            document.add_paragraph().add_run().add_break()


        # 5. Skills (from original extracted data, or could be a blend with suggested)
        if skills_list:
            all_skills = list(set(skills_list + suggested_skills))
            document.add_heading('Skills', level=2)
            document.add_paragraph(", ".join(all_skills))
            document.add_paragraph().add_run().add_break()

        # 6. Achievements (from original extracted data, if present)
        if achievements_list:
            document.add_heading('Achievements', level=2)
            for achievement in achievements_list:
                p = document.add_paragraph(style='List Bullet')
                p.add_run(achievement.strip())
            document.add_paragraph().add_run().add_break()


        byte_io = io.BytesIO()
        document.save(byte_io)
        byte_io.seek(0)

        return send_file(byte_io, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name='tailored_resume.docx')

    elif export_format == 'pdf':
        byte_io = io.BytesIO()
        # Corrected: Use 'inch' from reportlab.lib.units for margins
        doc = SimpleDocTemplate(byte_io, pagesize=letter,
                                leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                                topMargin=0.75 * inch, bottomMargin=0.75 * inch)
        styles = getSampleStyleSheet()
        
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['h1'],
            fontName='Helvetica-Bold',
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=6
        )
        contact_style = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        section_heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['h2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=16,
            alignment=TA_LEFT,
            spaceAfter=8,
            spaceBefore=12
        )
        body_text_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=6
        )
        list_style = styles['Bullet']

        story = []

        # 1. Contact Info (Header-like)
        story.append(Paragraph(name, name_style))
        contact_parts = []
        if email: contact_parts.append(email)
        if phone: contact_parts.append(phone)
        if linkedin: contact_parts.append(f"<link href='{linkedin}'>{linkedin}</link>")
        story.append(Paragraph(" | ".join(contact_parts), contact_style))
        story.append(Spacer(1, 0.15 * inch))


        # 2. Summary
        if improved_summary:
            story.append(Paragraph("SUMMARY", section_heading_style))
            story.append(Paragraph(improved_summary, body_text_style))
            story.append(Spacer(1, 0.15 * inch))

        # 3. Experience
        if improved_bullets:
            story.append(Paragraph("EXPERIENCE", section_heading_style))
            bullet_items = [ListItem(Paragraph(bullet.lstrip('* ').strip(), body_text_style)) for bullet in improved_bullets]
            story.append(ListFlowable(bullet_items, bulletType='bullet', start='bullet', leftIndent=0.2*inch))
            story.append(Spacer(1, 0.15 * inch))

        # 4. Education
        if education_entries:
            story.append(Paragraph("EDUCATION", section_heading_style))
            for edu in education_entries:
                edu_text = f"<b>{edu.get('degree', '')}</b>, {edu.get('university', '')}"
                if edu.get('year'):
                    edu_text += f", {edu.get('year', '')}"
                story.append(Paragraph(edu_text, body_text_style))
            story.append(Spacer(1, 0.15 * inch))

        # 5. Skills
        if skills_list:
            all_skills = list(set(skills_list + suggested_skills))
            if all_skills:
                story.append(Paragraph("SKILLS", section_heading_style))
                story.append(Paragraph(", ".join(all_skills), body_text_style))
                story.append(Spacer(1, 0.15 * inch))

        # 6. Achievements (from original extracted data, if present)
        if achievements_list:
            story.append(Paragraph("ACHIEVEMENTS", section_heading_style))
            bullet_items = [ListItem(Paragraph(ach.strip(), body_text_style)) for ach in achievements_list]
            story.append(ListFlowable(bullet_items, bulletType='bullet', start='bullet', leftIndent=0.2*inch))
            story.append(Spacer(1, 0.15 * inch))


        try:
            doc.build(story)
        except Exception as e:
            current_app.logger.error(f"Error building PDF: {str(e)}")
            return jsonify({"error": f"Error generating PDF: {str(e)}"}), 500

        byte_io.seek(0)

        return send_file(byte_io, mimetype='application/pdf',
                         as_attachment=True, download_name='tailored_resume.pdf')
    else:
        return jsonify({"error": "Unsupported export format."}), 400